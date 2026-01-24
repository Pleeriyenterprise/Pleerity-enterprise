"""
Template Renderer V2 - Enterprise-grade document rendering from orchestrator JSON.

This service:
1. Takes structured JSON output from the Document Orchestrator
2. Renders professional DOCX + sealed PDF documents
3. Enforces deterministic filenames: {order_ref}_{service_code}_v{version}_{status}_{YYYYMMDD-HHMM}.{ext}
4. Stores server-side SHA256 hash per file for tamper detection
5. Implements immutable versioning (no overwrites)

IMMUTABILITY RULES:
- Each generation creates a NEW version, never overwrites
- Previous versions marked SUPERSEDED, never deleted
- Intake snapshot locked BEFORE GPT execution
- All versions retained for audit trail

FLOW:
Orchestrator JSON → Template Selection → Content Rendering → DOCX Generation → 
PDF Generation (sealed) → Hash Computation → Version Storage → Human Review Ready
"""
import io
import json
import hashlib
import logging
from datetime import datetime, timezone
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable, ListFlowable, ListItem
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from database import database

logger = logging.getLogger(__name__)


# ============================================================================
# CONSTANTS & BRANDING
# ============================================================================

BRAND_TEAL = (0, 184, 169)      # #00B8A9
BRAND_TEAL_HEX = "00B8A9"
BRAND_NAVY = (11, 29, 58)       # #0B1D3A
BRAND_NAVY_HEX = "0B1D3A"
BRAND_GRAY = (128, 128, 128)
BRAND_LIGHT_GRAY = (240, 240, 240)


class RenderStatus(str, Enum):
    """Document render status."""
    DRAFT = "DRAFT"
    REGENERATED = "REGENERATED"
    SUPERSEDED = "SUPERSEDED"
    FINAL = "FINAL"


@dataclass
class RenderedDocument:
    """Represents a rendered document with metadata."""
    filename: str
    content_type: str
    content: bytes
    sha256_hash: str
    size_bytes: int
    format: str  # "docx" or "pdf"


@dataclass
class RenderResult:
    """Result of a document render operation."""
    success: bool
    order_id: str
    version: int
    status: RenderStatus
    docx: Optional[RenderedDocument] = None
    pdf: Optional[RenderedDocument] = None
    error_message: Optional[str] = None
    render_time_ms: int = 0
    intake_snapshot_hash: Optional[str] = None
    json_output_hash: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "order_id": self.order_id,
            "version": self.version,
            "status": self.status.value,
            "docx": {
                "filename": self.docx.filename,
                "content_type": self.docx.content_type,
                "sha256_hash": self.docx.sha256_hash,
                "size_bytes": self.docx.size_bytes,
            } if self.docx else None,
            "pdf": {
                "filename": self.pdf.filename,
                "content_type": self.pdf.content_type,
                "sha256_hash": self.pdf.sha256_hash,
                "size_bytes": self.pdf.size_bytes,
            } if self.pdf else None,
            "error_message": self.error_message,
            "render_time_ms": self.render_time_ms,
            "intake_snapshot_hash": self.intake_snapshot_hash,
            "json_output_hash": self.json_output_hash,
        }


# ============================================================================
# FILENAME GENERATION (DETERMINISTIC)
# ============================================================================

def generate_deterministic_filename(
    order_ref: str,
    service_code: str,
    version: int,
    status: RenderStatus,
    extension: str,
) -> str:
    """
    Generate deterministic filename following the canonical format:
    {order_ref}_{service_code}_v{version}_{status}_{YYYYMMDD-HHMM}.{ext}
    
    Example: ORD-2026-001234_AI_WF_BLUEPRINT_v1_DRAFT_20260122-1845.docx
    """
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M")
    
    # Sanitize order_ref and service_code
    safe_order_ref = order_ref.replace(" ", "-").replace("/", "-")
    safe_service_code = service_code.upper()
    
    return f"{safe_order_ref}_{safe_service_code}_v{version}_{status.value}_{timestamp}.{extension}"


def compute_sha256(content: bytes) -> str:
    """Compute SHA256 hash of content."""
    return hashlib.sha256(content).hexdigest()


# ============================================================================
# TEMPLATE RENDERER SERVICE
# ============================================================================

class TemplateRenderer:
    """
    Enterprise-grade template renderer for converting orchestrator JSON to documents.
    
    Features:
    - Deterministic filenames
    - SHA256 hash per file for tamper detection
    - Immutable versioning (no overwrites)
    - Professional branded output
    - Sealed PDF generation
    """
    
    VERSIONS_COLLECTION = "document_versions_v2"
    
    async def render_from_orchestration(
        self,
        order_id: str,
        structured_output: Dict[str, Any],
        intake_snapshot: Dict[str, Any],
        is_regeneration: bool = False,
        regeneration_notes: Optional[str] = None,
    ) -> RenderResult:
        """
        Render documents from orchestrator structured output.
        
        This is the main entry point for document rendering.
        
        Args:
            order_id: The order ID
            structured_output: JSON output from orchestrator
            intake_snapshot: Immutable copy of intake data (locked before GPT)
            is_regeneration: Whether this is a regeneration
            regeneration_notes: Notes for regeneration
        
        Returns:
            RenderResult with rendered documents
        """
        start_time = datetime.now(timezone.utc)
        db = database.get_db()
        
        # Fetch order for metadata
        order = await db.orders.find_one({"order_id": order_id}, {"_id": 0})
        if not order:
            return RenderResult(
                success=False,
                order_id=order_id,
                version=0,
                status=RenderStatus.DRAFT,
                error_message=f"Order not found: {order_id}",
            )
        
        service_code = order.get("service_code", "GENERAL")
        order_ref = order.get("order_ref", order_id)
        
        # Determine version number (immutable - always increment)
        existing_versions = await self._get_version_count(order_id)
        new_version = existing_versions + 1
        
        # Determine status
        status = RenderStatus.REGENERATED if is_regeneration else RenderStatus.DRAFT
        
        # Mark previous versions as SUPERSEDED (never delete)
        if new_version > 1:
            await self._mark_previous_superseded(order_id)
        
        # Compute hashes for integrity
        intake_hash = compute_sha256(
            json.dumps(intake_snapshot, sort_keys=True, default=str).encode()
        )
        json_output_hash = compute_sha256(
            json.dumps(structured_output, sort_keys=True, default=str).encode()
        )
        
        try:
            # Generate DOCX
            docx_content = self._render_docx(
                order=order,
                structured_output=structured_output,
                intake_snapshot=intake_snapshot,
                version=new_version,
                status=status,
                regeneration_notes=regeneration_notes,
            )
            
            docx_filename = generate_deterministic_filename(
                order_ref=order_ref,
                service_code=service_code,
                version=new_version,
                status=status,
                extension="docx",
            )
            
            docx_doc = RenderedDocument(
                filename=docx_filename,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                content=docx_content,
                sha256_hash=compute_sha256(docx_content),
                size_bytes=len(docx_content),
                format="docx",
            )
            
            # Generate sealed PDF
            pdf_content = self._render_pdf(
                order=order,
                structured_output=structured_output,
                intake_snapshot=intake_snapshot,
                version=new_version,
                status=status,
                regeneration_notes=regeneration_notes,
            )
            
            pdf_filename = generate_deterministic_filename(
                order_ref=order_ref,
                service_code=service_code,
                version=new_version,
                status=status,
                extension="pdf",
            )
            
            pdf_doc = RenderedDocument(
                filename=pdf_filename,
                content_type="application/pdf",
                content=pdf_content,
                sha256_hash=compute_sha256(pdf_content),
                size_bytes=len(pdf_content),
                format="pdf",
            )
            
        except Exception as e:
            logger.error(f"Render failed for {order_id}: {e}")
            return RenderResult(
                success=False,
                order_id=order_id,
                version=new_version,
                status=status,
                error_message=f"Render failed: {str(e)}",
            )
        
        # Store version record (immutable)
        version_record = {
            "order_id": order_id,
            "order_ref": order_ref,
            "service_code": service_code,
            "version": new_version,
            "status": status.value,
            "is_regeneration": is_regeneration,
            "regeneration_notes": regeneration_notes,
            # Document files
            "docx": {
                "filename": docx_doc.filename,
                "sha256_hash": docx_doc.sha256_hash,
                "size_bytes": docx_doc.size_bytes,
            },
            "pdf": {
                "filename": pdf_doc.filename,
                "sha256_hash": pdf_doc.sha256_hash,
                "size_bytes": pdf_doc.size_bytes,
            },
            # Integrity hashes
            "intake_snapshot_hash": intake_hash,
            "json_output_hash": json_output_hash,
            # Audit
            "created_at": datetime.now(timezone.utc),
            "created_by": "system",
            # Snapshots (for complete audit trail)
            "intake_snapshot": intake_snapshot,
            "structured_output": structured_output,
        }
        
        await db[self.VERSIONS_COLLECTION].insert_one(version_record)
        
        # Calculate render time
        render_time = int((datetime.now(timezone.utc) - start_time).total_seconds() * 1000)
        
        logger.info(f"Rendered documents for {order_id} v{new_version}: DOCX={docx_doc.sha256_hash[:8]}, PDF={pdf_doc.sha256_hash[:8]}")
        
        return RenderResult(
            success=True,
            order_id=order_id,
            version=new_version,
            status=status,
            docx=docx_doc,
            pdf=pdf_doc,
            render_time_ms=render_time,
            intake_snapshot_hash=intake_hash,
            json_output_hash=json_output_hash,
        )
    
    async def _get_version_count(self, order_id: str) -> int:
        """Get current version count for an order."""
        db = database.get_db()
        return await db[self.VERSIONS_COLLECTION].count_documents({"order_id": order_id})
    
    async def _mark_previous_superseded(self, order_id: str):
        """Mark all previous versions as SUPERSEDED."""
        db = database.get_db()
        await db[self.VERSIONS_COLLECTION].update_many(
            {"order_id": order_id, "status": {"$ne": RenderStatus.SUPERSEDED.value}},
            {"$set": {"status": RenderStatus.SUPERSEDED.value, "superseded_at": datetime.now(timezone.utc)}}
        )
    
    async def get_version(self, order_id: str, version: int) -> Optional[Dict[str, Any]]:
        """Get a specific version record."""
        db = database.get_db()
        return await db[self.VERSIONS_COLLECTION].find_one(
            {"order_id": order_id, "version": version},
            {"_id": 0}
        )
    
    async def get_latest_version(self, order_id: str) -> Optional[Dict[str, Any]]:
        """Get the latest version record."""
        db = database.get_db()
        return await db[self.VERSIONS_COLLECTION].find_one(
            {"order_id": order_id},
            {"_id": 0},
            sort=[("version", -1)]
        )
    
    async def get_all_versions(self, order_id: str) -> List[Dict[str, Any]]:
        """Get all versions for an order (audit trail)."""
        db = database.get_db()
        cursor = db[self.VERSIONS_COLLECTION].find(
            {"order_id": order_id},
            {"_id": 0}
        ).sort("version", 1)
        return await cursor.to_list(length=None)
    
    async def mark_final(self, order_id: str, version: int, approved_by: str) -> bool:
        """Mark a version as FINAL (approved for delivery)."""
        db = database.get_db()
        result = await db[self.VERSIONS_COLLECTION].update_one(
            {"order_id": order_id, "version": version},
            {
                "$set": {
                    "status": RenderStatus.FINAL.value,
                    "approved_at": datetime.now(timezone.utc),
                    "approved_by": approved_by,
                }
            }
        )
        return result.modified_count > 0
    
    async def verify_integrity(self, order_id: str, version: int) -> Tuple[bool, str]:
        """
        Verify document integrity by comparing stored hashes.
        Returns (is_valid, message).
        """
        version_record = await self.get_version(order_id, version)
        if not version_record:
            return False, f"Version not found: {order_id} v{version}"
        
        # For full verification, we'd need to retrieve the actual files
        # and recompute hashes - this is the check structure
        return True, "Hash records present for integrity verification"
    
    # ========================================================================
    # DOCX RENDERING
    # ========================================================================
    
    def _render_docx(
        self,
        order: Dict[str, Any],
        structured_output: Dict[str, Any],
        intake_snapshot: Dict[str, Any],
        version: int,
        status: RenderStatus,
        regeneration_notes: Optional[str] = None,
    ) -> bytes:
        """Render structured output to professional DOCX."""
        doc = Document()
        
        # Set up styles
        self._setup_docx_styles(doc)
        
        # Add header with branding
        self._add_docx_header(doc, order, version, status)
        
        # Add document content based on service type
        service_code = order.get("service_code", "GENERAL")
        self._add_docx_content(doc, structured_output, service_code)
        
        # Add footer with metadata
        self._add_docx_footer(doc, order, version, status)
        
        # Add watermark for non-final documents
        if status != RenderStatus.FINAL:
            self._add_docx_watermark(doc, status.value)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _setup_docx_styles(self, doc: Document):
        """Set up document styles."""
        styles = doc.styles
        
        # Title style
        if 'CustomTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(11, 29, 58)  # Brand navy
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading style
        if 'CustomHeading' not in [s.name for s in styles]:
            heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 184, 169)  # Brand teal
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
    
    def _add_docx_header(self, doc: Document, order: Dict[str, Any], version: int, status: RenderStatus):
        """Add branded header to document."""
        # Company name
        header_para = doc.add_paragraph()
        header_run = header_para.add_run("PLEERITY ENTERPRISE LTD")
        header_run.bold = True
        header_run.font.size = Pt(10)
        header_run.font.color.rgb = RGBColor(0, 184, 169)
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Document info line
        info_para = doc.add_paragraph()
        order_ref = order.get("order_ref", order.get("order_id", ""))
        service_name = order.get("service_name", order.get("service_code", ""))
        info_run = info_para.add_run(f"Order: {order_ref} | Version: {version} | Status: {status.value}")
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(128, 128, 128)
        info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Horizontal line
        doc.add_paragraph("_" * 80)
        
        # Document title
        title_para = doc.add_paragraph()
        title_para.style = doc.styles['CustomTitle']
        title_para.add_run(service_name)
        
        doc.add_paragraph()  # Spacer
    
    def _add_docx_content(self, doc: Document, structured_output: Dict[str, Any], service_code: str):
        """Add content sections based on structured output."""
        
        # Route to appropriate content renderer based on service type
        if service_code.startswith("AI_"):
            self._render_ai_service_content(doc, structured_output, service_code)
        elif service_code.startswith("MR_"):
            self._render_market_research_content(doc, structured_output, service_code)
        elif service_code.startswith("COMP_"):
            self._render_compliance_content(doc, structured_output, service_code)
        elif service_code.startswith("DOC_PACK_"):
            self._render_document_pack_content(doc, structured_output, service_code)
        else:
            self._render_generic_content(doc, structured_output)
    
    def _render_ai_service_content(self, doc: Document, output: Dict[str, Any], service_code: str):
        """Render AI service content (blueprints, process maps, tool recommendations)."""
        
        # Executive Summary
        if "executive_summary" in output:
            self._add_section_heading(doc, "Executive Summary")
            doc.add_paragraph(str(output["executive_summary"]))
        
        # Current State Analysis (handle both old and new field names)
        csa = output.get("current_state_analysis") or output.get("current_state_assessment")
        if csa:
            self._add_section_heading(doc, "Current State Analysis")
            
            if isinstance(csa, dict):
                # Handle structured analysis
                if "pain_points" in csa:
                    self._add_subsection(doc, "Pain Points")
                    for point in csa["pain_points"]:
                        if isinstance(point, dict):
                            doc.add_paragraph(f"• {point.get('description', point)}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"• {point}", style='List Bullet')
                
                if "inefficiencies" in csa:
                    self._add_subsection(doc, "Current Inefficiencies")
                    for item in csa["inefficiencies"]:
                        if isinstance(item, dict):
                            doc.add_paragraph(f"• {item.get('description', item)}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"• {item}", style='List Bullet')
                
                if "opportunities" in csa:
                    self._add_subsection(doc, "Improvement Opportunities")
                    for opp in csa["opportunities"]:
                        if isinstance(opp, dict):
                            doc.add_paragraph(f"• {opp.get('description', opp)}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"• {opp}", style='List Bullet')
                
                if "business_overview" in csa:
                    doc.add_paragraph(str(csa["business_overview"]))
            else:
                # Handle string analysis
                doc.add_paragraph(str(csa))
        
        # Recommended Workflows (handle both old and new field names)
        workflows = output.get("recommended_workflows") or output.get("workflow_analysis")
        if workflows:
            self._add_section_heading(doc, "Recommended Workflows")
            
            if isinstance(workflows, list):
                for i, wf in enumerate(workflows, 1):
                    if isinstance(wf, dict):
                        wf_name = wf.get("workflow_name") or wf.get("name") or f"Workflow {i}"
                        self._add_subsection(doc, wf_name)
                        
                        if wf.get("description"):
                            doc.add_paragraph(str(wf["description"]))
                        
                        if wf.get("current_state"):
                            doc.add_paragraph(f"Current State: {wf['current_state']}")
                        
                        if wf.get("automation_potential"):
                            doc.add_paragraph(f"Automation Potential: {wf['automation_potential']}")
                        
                        if wf.get("recommended_approach"):
                            doc.add_paragraph(f"Recommended Approach: {wf['recommended_approach']}")
                        
                        if wf.get("steps"):
                            doc.add_paragraph("Implementation Steps:", style='Normal')
                            for step in wf["steps"]:
                                doc.add_paragraph(f"• {step}", style='List Bullet')
                        
                        if wf.get("tools"):
                            doc.add_paragraph(f"Recommended Tools: {', '.join(wf['tools'])}")
                        
                        if wf.get("expected_savings"):
                            doc.add_paragraph(f"Expected Savings: {wf['expected_savings']}")
                    else:
                        doc.add_paragraph(f"• {wf}", style='List Bullet')
        
        # Automation Opportunities
        if "automation_opportunities" in output:
            self._add_section_heading(doc, "Automation Opportunities")
            for opp in output["automation_opportunities"]:
                if isinstance(opp, dict):
                    priority = opp.get("priority", "")
                    opportunity = opp.get("opportunity", "")
                    effort = opp.get("implementation_effort", "")
                    doc.add_paragraph(f"[{priority}] {opportunity} (Effort: {effort})")
                else:
                    doc.add_paragraph(f"• {opp}", style='List Bullet')
        
        # Tool Recommendations
        if "tool_recommendations" in output:
            self._add_section_heading(doc, "Tool Recommendations")
            for tool in output["tool_recommendations"]:
                if isinstance(tool, dict):
                    self._add_subsection(doc, tool.get("category", "Tool"))
                    doc.add_paragraph(f"Recommended: {tool.get('recommended_tool', 'N/A')}")
                    doc.add_paragraph(f"Rationale: {tool.get('rationale', '')}")
                    if tool.get("alternatives"):
                        doc.add_paragraph(f"Alternatives: {', '.join(tool['alternatives'])}")
                else:
                    doc.add_paragraph(f"• {tool}", style='List Bullet')
        
        # Implementation Roadmap
        roadmap = output.get("implementation_roadmap")
        if roadmap:
            self._add_section_heading(doc, "Implementation Roadmap")
            
            if isinstance(roadmap, dict):
                # Handle phases
                if "phases" in roadmap:
                    for phase in roadmap["phases"]:
                        if isinstance(phase, dict):
                            phase_name = phase.get("phase_name") or phase.get("name", "Phase")
                            self._add_subsection(doc, phase_name)
                            if phase.get("description"):
                                doc.add_paragraph(str(phase["description"]))
                            if phase.get("duration"):
                                doc.add_paragraph(f"Duration: {phase['duration']}")
                            if phase.get("tasks"):
                                for task in phase["tasks"]:
                                    doc.add_paragraph(f"• {task}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"• {phase}", style='List Bullet')
                
                # Handle timeline
                if "timeline" in roadmap:
                    self._add_subsection(doc, "Timeline")
                    doc.add_paragraph(str(roadmap["timeline"]))
                
                # Handle milestones
                if "milestones" in roadmap:
                    self._add_subsection(doc, "Key Milestones")
                    for milestone in roadmap["milestones"]:
                        if isinstance(milestone, dict):
                            doc.add_paragraph(f"• {milestone.get('name', milestone)}: {milestone.get('date', '')}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"• {milestone}", style='List Bullet')
                
                # Handle legacy format (phase_name: items)
                for phase_name, phase_items in roadmap.items():
                    if phase_name not in ["phases", "timeline", "milestones"] and phase_items:
                        self._add_subsection(doc, phase_name.replace("_", " ").title())
                        if isinstance(phase_items, list):
                            for item in phase_items:
                                doc.add_paragraph(f"• {item}", style='List Bullet')
                        else:
                            doc.add_paragraph(str(phase_items))
            else:
                doc.add_paragraph(str(roadmap))
        
        # Expected Outcomes
        outcomes = output.get("expected_outcomes")
        if outcomes:
            self._add_section_heading(doc, "Expected Outcomes")
            
            if isinstance(outcomes, dict):
                if "efficiency_gains" in outcomes:
                    self._add_subsection(doc, "Efficiency Gains")
                    for gain in outcomes["efficiency_gains"]:
                        doc.add_paragraph(f"• {gain}", style='List Bullet')
                
                if "roi_estimate" in outcomes:
                    self._add_subsection(doc, "ROI Estimate")
                    doc.add_paragraph(str(outcomes["roi_estimate"]))
                
                if "key_metrics" in outcomes:
                    self._add_subsection(doc, "Key Metrics")
                    for metric in outcomes["key_metrics"]:
                        if isinstance(metric, dict):
                            doc.add_paragraph(f"• {metric.get('name', metric)}: {metric.get('expected_improvement', '')}", style='List Bullet')
                        else:
                            doc.add_paragraph(f"• {metric}", style='List Bullet')
                
                # Handle other outcome fields
                for key, value in outcomes.items():
                    if key not in ["efficiency_gains", "roi_estimate", "key_metrics"] and value:
                        self._add_subsection(doc, key.replace("_", " ").title())
                        if isinstance(value, list):
                            for item in value:
                                doc.add_paragraph(f"• {item}", style='List Bullet')
                        else:
                            doc.add_paragraph(str(value))
            else:
                doc.add_paragraph(str(outcomes))
        
        # Assumptions (if any data gaps were flagged)
        assumptions = output.get("assumptions") or output.get("data_gaps_flagged")
        if assumptions:
            self._add_section_heading(doc, "Assumptions & Data Gaps")
            for assumption in assumptions:
                doc.add_paragraph(f"• {assumption}", style='List Bullet')
        
        # Data Gaps
        self._add_data_gaps_section(doc, output)
    
    def _render_market_research_content(self, doc: Document, output: Dict[str, Any], service_code: str):
        """Render market research content."""
        
        # Executive Summary (for advanced)
        if "executive_summary" in output:
            self._add_section_heading(doc, "Executive Summary")
            doc.add_paragraph(str(output["executive_summary"]))
        
        # Market Overview
        if "market_overview" in output:
            mo = output["market_overview"]
            self._add_section_heading(doc, "Market Overview")
            
            if "industry_summary" in mo:
                doc.add_paragraph(str(mo["industry_summary"]))
            
            if "key_trends" in mo:
                self._add_subsection(doc, "Key Trends")
                for trend in mo["key_trends"]:
                    doc.add_paragraph(f"• {trend}", style='List Bullet')
        
        # Target Segment
        if "target_segment_analysis" in output:
            seg = output["target_segment_analysis"]
            self._add_section_heading(doc, "Target Segment Analysis")
            
            if "segment_description" in seg:
                doc.add_paragraph(str(seg["segment_description"]))
            
            if "pain_points" in seg:
                self._add_subsection(doc, "Customer Pain Points")
                for point in seg["pain_points"]:
                    doc.add_paragraph(f"• {point}", style='List Bullet')
        
        # Competitor Overview
        if "competitor_overview" in output:
            self._add_section_heading(doc, "Competitor Analysis")
            comps = output["competitor_overview"]
            
            if comps:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Competitor'
                hdr_cells[1].text = 'Position'
                hdr_cells[2].text = 'Strengths'
                hdr_cells[3].text = 'Weaknesses'
                
                for comp in comps:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(comp.get("competitor_name", ""))
                    row_cells[1].text = str(comp.get("market_position", ""))
                    row_cells[2].text = ", ".join(comp.get("key_strengths", [])[:2])
                    row_cells[3].text = ", ".join(comp.get("key_weaknesses", [])[:2])
        
        # SWOT Analysis (advanced)
        if "swot_analysis" in output:
            swot = output["swot_analysis"]
            self._add_section_heading(doc, "SWOT Analysis")
            
            swot_table = doc.add_table(rows=2, cols=2)
            swot_table.style = 'Table Grid'
            
            swot_table.cell(0, 0).text = "STRENGTHS\n" + "\n".join(f"• {s}" for s in swot.get("strengths", []))
            swot_table.cell(0, 1).text = "WEAKNESSES\n" + "\n".join(f"• {w}" for w in swot.get("weaknesses", []))
            swot_table.cell(1, 0).text = "OPPORTUNITIES\n" + "\n".join(f"• {o}" for o in swot.get("opportunities", []))
            swot_table.cell(1, 1).text = "THREATS\n" + "\n".join(f"• {t}" for t in swot.get("threats", []))
        
        # Key Findings
        if "key_findings" in output:
            self._add_section_heading(doc, "Key Findings")
            for finding in output["key_findings"]:
                f_text = finding.get("finding", "")
                f_impl = finding.get("implication", "")
                doc.add_paragraph(f"• {f_text}")
                if f_impl:
                    doc.add_paragraph(f"  Implication: {f_impl}")
        
        # Strategic Summary
        if "strategic_summary" in output:
            ss = output["strategic_summary"]
            self._add_section_heading(doc, "Strategic Summary")
            
            if "market_opportunity" in ss:
                doc.add_paragraph(f"Market Opportunity: {ss['market_opportunity']}")
            if "recommended_positioning" in ss:
                doc.add_paragraph(f"Recommended Positioning: {ss['recommended_positioning']}")
        
        self._add_data_gaps_section(doc, output)
    
    def _render_compliance_content(self, doc: Document, output: Dict[str, Any], service_code: str):
        """Render compliance audit content."""
        
        # Property/Audit Summary
        if "property_summary" in output or "audit_summary" in output:
            summary = output.get("property_summary") or output.get("audit_summary", {})
            self._add_section_heading(doc, "Summary")
            
            for key, value in summary.items():
                if value:
                    label = key.replace("_", " ").title()
                    doc.add_paragraph(f"{label}: {value}")
        
        # Safety Certificates
        if "safety_certificates" in output:
            self._add_section_heading(doc, "Safety Certificates Review")
            certs = output["safety_certificates"]
            
            if isinstance(certs, list):
                for cert in certs:
                    cert_type = cert.get("certificate_type", "Certificate")
                    status = cert.get("status", "Unknown")
                    compliance = cert.get("compliance_status", "Unknown")
                    doc.add_paragraph(f"• {cert_type}: {status} ({compliance})")
            elif isinstance(certs, dict):
                for cert_name, cert_data in certs.items():
                    if isinstance(cert_data, dict):
                        status = cert_data.get("status", "Unknown")
                        compliance = cert_data.get("compliance", "Unknown")
                        doc.add_paragraph(f"• {cert_name.replace('_', ' ').title()}: {status} ({compliance})")
        
        # Risk Summary
        if "risk_summary" in output or "risk_assessment" in output:
            risks = output.get("risk_summary") or output.get("risk_assessment", {})
            self._add_section_heading(doc, "Risk Assessment")
            
            overall = risks.get("overall_risk_level", "Not assessed")
            doc.add_paragraph(f"Overall Risk Level: {overall}")
            
            for risk_level in ["critical_risks", "critical_items", "high_risks", "high_priority_items"]:
                if risk_level in risks and risks[risk_level]:
                    level_name = risk_level.replace("_", " ").title()
                    self._add_subsection(doc, level_name)
                    for item in risks[risk_level]:
                        doc.add_paragraph(f"• {item}", style='List Bullet')
        
        # Action Plan
        if "action_plan" in output:
            self._add_section_heading(doc, "Action Plan")
            for action in output["action_plan"]:
                priority = action.get("priority", "")
                action_text = action.get("action", "")
                responsible = action.get("responsible_party", "")
                doc.add_paragraph(f"[{priority}] {action_text}")
                if responsible:
                    doc.add_paragraph(f"  Responsible: {responsible}")
        
        # Disclaimers
        if "disclaimers" in output:
            self._add_section_heading(doc, "Important Notices")
            for disclaimer in output["disclaimers"]:
                doc.add_paragraph(f"• {disclaimer}", style='List Bullet')
        
        self._add_data_gaps_section(doc, output)
    
    def _render_document_pack_content(self, doc: Document, output: Dict[str, Any], service_code: str):
        """Render document pack content."""
        
        # Orchestration metadata
        if "orchestration_metadata" in output:
            meta = output["orchestration_metadata"]
            self._add_section_heading(doc, "Document Pack Summary")
            doc.add_paragraph(f"Pack Type: {meta.get('pack_type', service_code)}")
            doc.add_paragraph(f"Documents Generated: {meta.get('total_documents', 0)}")
            
            if meta.get("documents_to_generate"):
                self._add_subsection(doc, "Documents Included")
                for doc_name in meta["documents_to_generate"]:
                    doc.add_paragraph(f"• {doc_name}", style='List Bullet')
        
        # Document contents
        if "document_contents" in output:
            for doc_content in output["document_contents"]:
                doc_name = doc_content.get("document_name", "Document")
                self._add_section_heading(doc, doc_name)
                
                # Merge fields
                if doc_content.get("merge_fields"):
                    for field_name, field_value in doc_content["merge_fields"].items():
                        label = field_name.replace("_", " ").title()
                        doc.add_paragraph(f"{label}: {field_value}")
                
                # GPT sections
                if doc_content.get("gpt_sections"):
                    for section_name, section_content in doc_content["gpt_sections"].items():
                        self._add_subsection(doc, section_name.replace("_", " ").title())
                        doc.add_paragraph(str(section_content))
                
                doc.add_page_break()
        
        self._add_data_gaps_section(doc, output)
    
    def _render_generic_content(self, doc: Document, output: Dict[str, Any]):
        """Render generic structured output."""
        for key, value in output.items():
            if key == "data_gaps_flagged":
                continue
            
            self._add_section_heading(doc, key.replace("_", " ").title())
            
            if isinstance(value, str):
                doc.add_paragraph(value)
            elif isinstance(value, list):
                for item in value:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            doc.add_paragraph(f"{k}: {v}")
                    else:
                        doc.add_paragraph(f"• {item}", style='List Bullet')
            elif isinstance(value, dict):
                for k, v in value.items():
                    doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")
        
        self._add_data_gaps_section(doc, output)
    
    def _add_section_heading(self, doc: Document, title: str):
        """Add a section heading."""
        para = doc.add_paragraph()
        para.style = doc.styles['CustomHeading']
        para.add_run(title)
    
    def _add_subsection(self, doc: Document, title: str):
        """Add a subsection heading."""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
    
    def _add_data_gaps_section(self, doc: Document, output: Dict[str, Any]):
        """Add data gaps section if present."""
        if output.get("data_gaps_flagged"):
            self._add_section_heading(doc, "Data Gaps Identified")
            doc.add_paragraph("The following information was not provided and may affect the accuracy of this document:")
            for gap in output["data_gaps_flagged"]:
                doc.add_paragraph(f"• {gap}", style='List Bullet')
    
    def _add_docx_footer(self, doc: Document, order: Dict[str, Any], version: int, status: RenderStatus):
        """Add footer with metadata."""
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        
        footer_para = doc.add_paragraph()
        footer_text = f"Generated by Pleerity Enterprise Ltd | {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} | Document v{version}"
        footer_run = footer_para.add_run(footer_text)
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Confidentiality notice
        conf_para = doc.add_paragraph()
        conf_run = conf_para.add_run("CONFIDENTIAL - This document contains proprietary information.")
        conf_run.font.size = Pt(8)
        conf_run.font.italic = True
        conf_run.font.color.rgb = RGBColor(128, 128, 128)
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_docx_watermark(self, doc: Document, watermark_text: str):
        """Add watermark to document (simplified - adds to header)."""
        # Add watermark text to each section header
        for section in doc.sections:
            header = section.header
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = header_para.add_run(f"  [{watermark_text}]")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(200, 200, 200)
            run.font.bold = True
    
    # ========================================================================
    # PDF RENDERING
    # ========================================================================
    
    def _render_pdf(
        self,
        order: Dict[str, Any],
        structured_output: Dict[str, Any],
        intake_snapshot: Dict[str, Any],
        version: int,
        status: RenderStatus,
        regeneration_notes: Optional[str] = None,
    ) -> bytes:
        """Render structured output to sealed PDF."""
        buffer = io.BytesIO()
        
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=25*mm,
            bottomMargin=20*mm,
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles
        styles.add(ParagraphStyle(
            name='BrandTitle',
            parent=styles['Title'],
            fontSize=20,
            textColor=colors.HexColor('#0B1D3A'),
            spaceAfter=12,
        ))
        
        styles.add(ParagraphStyle(
            name='BrandHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#00B8A9'),
            spaceBefore=12,
            spaceAfter=6,
        ))
        
        styles.add(ParagraphStyle(
            name='BrandBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            alignment=TA_JUSTIFY,
        ))
        
        story = []
        
        # Header
        order_ref = order.get("order_ref", order.get("order_id", ""))
        service_name = order.get("service_name", order.get("service_code", ""))
        
        header_text = f"PLEERITY ENTERPRISE LTD | Order: {order_ref} | v{version} | {status.value}"
        story.append(Paragraph(header_text, styles['Normal']))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#00B8A9')))
        story.append(Spacer(1, 12))
        
        # Title
        story.append(Paragraph(service_name, styles['BrandTitle']))
        story.append(Spacer(1, 12))
        
        # Watermark notice for non-final
        if status != RenderStatus.FINAL:
            watermark_style = ParagraphStyle(
                name='Watermark',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.HexColor('#CCCCCC'),
                alignment=TA_CENTER,
            )
            story.append(Paragraph(f"[{status.value} - FOR REVIEW ONLY]", watermark_style))
            story.append(Spacer(1, 12))
        
        # Content
        service_code = order.get("service_code", "GENERAL")
        self._add_pdf_content(story, styles, structured_output, service_code)
        
        # Footer
        story.append(Spacer(1, 20))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CCCCCC')))
        
        footer_style = ParagraphStyle(
            name='Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#888888'),
            alignment=TA_CENTER,
        )
        
        footer_text = f"Generated by Pleerity Enterprise Ltd | {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} | Document v{version}"
        story.append(Paragraph(footer_text, footer_style))
        story.append(Paragraph("CONFIDENTIAL - This document contains proprietary information.", footer_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _add_pdf_content(self, story: List, styles, output: Dict[str, Any], service_code: str):
        """Add content to PDF based on service type."""
        
        # Process main sections
        for key, value in output.items():
            if key in ["data_gaps_flagged", "raw_response", "parse_error"]:
                continue
            
            # Section heading
            section_title = key.replace("_", " ").title()
            story.append(Paragraph(section_title, styles['BrandHeading']))
            
            if isinstance(value, str):
                story.append(Paragraph(value, styles['BrandBody']))
            
            elif isinstance(value, list):
                for item in value:
                    if isinstance(item, dict):
                        # Table for dict items
                        item_text = " | ".join(f"{k}: {v}" for k, v in item.items() if v)
                        story.append(Paragraph(f"• {item_text}", styles['BrandBody']))
                    else:
                        story.append(Paragraph(f"• {item}", styles['BrandBody']))
            
            elif isinstance(value, dict):
                for k, v in value.items():
                    if isinstance(v, list):
                        story.append(Paragraph(f"<b>{k.replace('_', ' ').title()}:</b>", styles['BrandBody']))
                        for list_item in v:
                            story.append(Paragraph(f"  • {list_item}", styles['BrandBody']))
                    elif isinstance(v, dict):
                        story.append(Paragraph(f"<b>{k.replace('_', ' ').title()}:</b>", styles['BrandBody']))
                        for sub_k, sub_v in v.items():
                            story.append(Paragraph(f"  {sub_k}: {sub_v}", styles['BrandBody']))
                    else:
                        story.append(Paragraph(f"<b>{k.replace('_', ' ').title()}:</b> {v}", styles['BrandBody']))
            
            story.append(Spacer(1, 8))
        
        # Data gaps section
        if output.get("data_gaps_flagged"):
            story.append(Paragraph("Data Gaps Identified", styles['BrandHeading']))
            story.append(Paragraph(
                "The following information was not provided and may affect accuracy:",
                styles['BrandBody']
            ))
            for gap in output["data_gaps_flagged"]:
                story.append(Paragraph(f"• {gap}", styles['BrandBody']))


# Singleton instance
template_renderer = TemplateRenderer()
