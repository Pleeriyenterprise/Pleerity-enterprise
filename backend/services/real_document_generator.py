"""
Real Document Generator - Production-ready PDF/DOCX generation.
Uses reportlab for PDF and python-docx for DOCX files.

Key Features:
- Generates professional, branded PDF and DOCX documents
- Input data snapshotting for perfect regeneration
- Proper document formatting with headers, footers, tables
- Watermark support for draft documents
- Version tracking with status labels
"""
import io
import json
import hashlib
import logging
from datetime import datetime, timezone
from typing import Dict, Any, List, Optional

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm, inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, Image as RLImage, HRFlowable
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from services.storage_adapter import upload_order_document
from services.document_generator import (
    DocumentGenerator, DocumentVersion, DocumentStatus, DocumentType,
    generate_document_filename
)
from database import database

logger = logging.getLogger(__name__)


# Brand colors
BRAND_TEAL = (0, 184, 169)  # #00B8A9
BRAND_NAVY = (11, 29, 58)   # #0B1D3A
BRAND_GRAY = (128, 128, 128)


class RealDocumentGenerator(DocumentGenerator):
    """
    Production document generator using reportlab and python-docx.
    Generates professional PDF and DOCX files with proper branding.
    """
    
    SERVICE_TO_DOC_TYPE = {
        "DOC_PACK_TENANCY": DocumentType.SECTION_21_NOTICE,
        "DOC_PACK_ESSENTIAL": DocumentType.TENANCY_AGREEMENT,
        "DOC_PACK_ULTIMATE": DocumentType.TENANCY_AGREEMENT,
        "DOC_PACK_INVENTORY": DocumentType.INVENTORY_REPORT,
        "HMO_AUDIT": DocumentType.COMPLIANCE_AUDIT,
        "FULL_AUDIT": DocumentType.COMPLIANCE_AUDIT,
        "MR_BASIC": DocumentType.MARKET_RESEARCH,
        "MR_ADV": DocumentType.MARKET_RESEARCH,
        "AI_WF_BLUEPRINT": DocumentType.GENERAL_DOCUMENT,
        "AI_PROC_MAP": DocumentType.GENERAL_DOCUMENT,
        "AI_TOOL_REC": DocumentType.GENERAL_DOCUMENT,
        "MOVE_IN_OUT": DocumentType.INVENTORY_REPORT,
    }
    
    def get_document_type_for_service(self, service_code: str) -> DocumentType:
        return self.SERVICE_TO_DOC_TYPE.get(service_code, DocumentType.GENERAL_DOCUMENT)
    
    async def generate_documents(
        self,
        order_id: str,
        regeneration_notes: Optional[str] = None,
        regenerated_from_version: Optional[int] = None,
    ) -> DocumentVersion:
        """
        Generate real DOCX + PDF for an order with input data snapshot.
        """
        db = database.get_db()
        
        # Fetch order details
        order = await db.orders.find_one({"order_id": order_id})
        if not order:
            raise ValueError(f"Order not found: {order_id}")
        
        # Determine version number
        existing_versions = order.get("document_versions", [])
        new_version = len(existing_versions) + 1
        
        # Get service code
        service_code = order.get("service_code", "GENERAL")
        
        # Determine status
        is_regeneration = new_version > 1 or regeneration_notes is not None
        status = DocumentStatus.REGENERATED if is_regeneration else DocumentStatus.DRAFT
        
        # Mark previous versions as SUPERSEDED
        if new_version > 1:
            await self._mark_previous_versions_superseded(db, order_id)
        
        # Get document type
        doc_type = self.get_document_type_for_service(service_code)
        
        # Create input data snapshot for traceability and regeneration
        input_snapshot = self._create_input_snapshot(order)
        input_data_hash = hashlib.sha256(
            json.dumps(input_snapshot, sort_keys=True, default=str).encode()
        ).hexdigest()
        
        # Generate DOCX
        docx_buffer = self._generate_docx(order, new_version, doc_type, status, regeneration_notes)
        
        # Generate PDF
        pdf_buffer = self._generate_pdf(order, new_version, doc_type, status, regeneration_notes)
        
        # Generate proper filenames
        docx_filename = generate_document_filename(
            order_id=order_id,
            service_code=service_code,
            version=new_version,
            status=status,
            extension="docx",
        )
        pdf_filename = generate_document_filename(
            order_id=order_id,
            service_code=service_code,
            version=new_version,
            status=status,
            extension="pdf",
        )
        
        # Upload DOCX to storage
        docx_meta = await upload_order_document(
            order_id=order_id,
            file_data=io.BytesIO(docx_buffer.getvalue()),
            filename=docx_filename,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            document_type=doc_type.value,
            version=new_version,
            uploaded_by="system",
        )
        
        # Upload PDF to storage
        pdf_meta = await upload_order_document(
            order_id=order_id,
            file_data=io.BytesIO(pdf_buffer.getvalue()),
            filename=pdf_filename,
            content_type="application/pdf",
            document_type=doc_type.value,
            version=new_version,
            uploaded_by="system",
        )
        
        # Create version record with input snapshot
        doc_version = DocumentVersion(
            version=new_version,
            document_type=doc_type,
            status=status,
            file_id_docx=docx_meta.file_id,
            file_id_pdf=pdf_meta.file_id,
            filename_docx=docx_filename,
            filename_pdf=pdf_filename,
            generated_at=datetime.now(timezone.utc),
            generated_by="system",
            is_regeneration=is_regeneration,
            regeneration_notes=regeneration_notes,
            regenerated_from_version=regenerated_from_version,
            content_hash=pdf_meta.sha256_hash,
            input_data_hash=input_data_hash,
        )
        
        # Update order with new version and input snapshot
        version_dict = doc_version.to_dict()
        version_dict["input_snapshot"] = input_snapshot  # Store snapshot for regeneration
        
        await db.orders.update_one(
            {"order_id": order_id},
            {
                "$push": {"document_versions": version_dict},
                "$set": {
                    "current_document_version": new_version,
                    "updated_at": datetime.now(timezone.utc),
                }
            }
        )
        
        logger.info(f"Generated real documents for order {order_id} v{new_version}")
        return doc_version
    
    def _create_input_snapshot(self, order: Dict[str, Any]) -> Dict[str, Any]:
        """
        Create a snapshot of all input data used for document generation.
        This allows perfect regeneration with the same inputs.
        """
        return {
            "snapshot_created_at": datetime.now(timezone.utc).isoformat(),
            "order_id": order.get("order_id"),
            "order_ref": order.get("order_ref"),
            "service_code": order.get("service_code"),
            "service_name": order.get("service_name"),
            "customer": {
                "full_name": order.get("customer", {}).get("full_name"),
                "email": order.get("customer", {}).get("email"),
                "phone": order.get("customer", {}).get("phone"),
                "company": order.get("customer", {}).get("company"),
            },
            "parameters": order.get("parameters", {}),
            "client_input_responses": order.get("client_input_response", []),
        }
    
    async def _mark_previous_versions_superseded(self, db, order_id: str):
        """Mark all previous non-final versions as SUPERSEDED."""
        await db.orders.update_one(
            {"order_id": order_id},
            {
                "$set": {
                    "document_versions.$[elem].status": DocumentStatus.SUPERSEDED.value
                }
            },
            array_filters=[
                {"elem.status": {"$nin": [DocumentStatus.FINAL.value, DocumentStatus.VOID.value]}}
            ]
        )
    
    # =========================================================================
    # DOCX GENERATION
    # =========================================================================
    
    def _generate_docx(
        self,
        order: Dict[str, Any],
        version: int,
        doc_type: DocumentType,
        status: DocumentStatus,
        regeneration_notes: Optional[str] = None,
    ) -> io.BytesIO:
        """Generate a professional DOCX document."""
        doc = Document()
        
        # Set document properties
        core_props = doc.core_properties
        core_props.author = "Pleerity Enterprise Ltd"
        core_props.title = f"{order.get('service_name', 'Document')} - {order.get('order_id')}"
        core_props.subject = doc_type.value
        
        # Add styles
        self._setup_docx_styles(doc)
        
        # Add header with status watermark if DRAFT
        self._add_docx_header(doc, order, version, status)
        
        # Add main content based on document type
        self._add_docx_content(doc, order, version, doc_type, regeneration_notes)
        
        # Add footer
        self._add_docx_footer(doc, order, version, status)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _setup_docx_styles(self, doc: Document):
        """Setup custom styles for the document."""
        styles = doc.styles
        
        # Title style
        if 'CustomTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(*BRAND_NAVY)
            title_style.paragraph_format.space_after = Pt(12)
        
        # Section heading style
        if 'SectionHeading' not in [s.name for s in styles]:
            section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(*BRAND_TEAL)
            section_style.paragraph_format.space_before = Pt(12)
            section_style.paragraph_format.space_after = Pt(6)
    
    def _add_docx_header(
        self,
        doc: Document,
        order: Dict[str, Any],
        version: int,
        status: DocumentStatus,
    ):
        """Add document header with branding and status."""
        # Company header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run("PLEERITY ENTERPRISE LTD")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(*BRAND_NAVY)
        
        # Status watermark for drafts
        if status != DocumentStatus.FINAL:
            watermark_para = doc.add_paragraph()
            watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = watermark_para.add_run(f"*** {status.value} DOCUMENT ***")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 0, 0)
        
        doc.add_paragraph()  # Spacer
        
        # Document title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(order.get("service_name", "Document"))
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(*BRAND_NAVY)
        
        # Order reference and version
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_para.add_run(f"Reference: {order.get('order_id')} | Version: v{version}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*BRAND_GRAY)
        
        # Horizontal line
        doc.add_paragraph("_" * 80)
        doc.add_paragraph()
    
    def _add_docx_content(
        self,
        doc: Document,
        order: Dict[str, Any],
        version: int,
        doc_type: DocumentType,
        regeneration_notes: Optional[str] = None,
    ):
        """Add main document content based on type."""
        customer = order.get("customer", {})
        params = order.get("parameters", {})
        
        # Customer Details section
        self._add_docx_section(doc, "CUSTOMER DETAILS")
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        data = [
            ("Full Name:", customer.get("full_name", "Not provided")),
            ("Email:", customer.get("email", "Not provided")),
            ("Phone:", customer.get("phone", "Not provided")),
            ("Company:", customer.get("company", "N/A")),
        ]
        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Service Details section
        self._add_docx_section(doc, "SERVICE DETAILS")
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        service_data = [
            ("Service:", order.get("service_name", "N/A")),
            ("Service Code:", order.get("service_code", "N/A")),
            ("Category:", order.get("service_category", "N/A")),
        ]
        for i, (label, value) in enumerate(service_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()
        
        # Document-type specific content
        if doc_type == DocumentType.SECTION_21_NOTICE:
            self._add_section_21_content_docx(doc, params)
        elif doc_type == DocumentType.TENANCY_AGREEMENT:
            self._add_tenancy_content_docx(doc, params)
        elif doc_type == DocumentType.COMPLIANCE_AUDIT:
            self._add_audit_content_docx(doc, params)
        elif doc_type == DocumentType.MARKET_RESEARCH:
            self._add_market_research_content_docx(doc, params)
        elif doc_type == DocumentType.INVENTORY_REPORT:
            self._add_inventory_content_docx(doc, params)
        else:
            self._add_general_content_docx(doc, params)
        
        # Regeneration notes if present
        if regeneration_notes:
            doc.add_paragraph()
            self._add_docx_section(doc, "REGENERATION NOTES")
            notes_para = doc.add_paragraph(regeneration_notes)
            notes_para.style = 'Quote'
    
    def _add_docx_section(self, doc: Document, title: str):
        """Add a section heading."""
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*BRAND_TEAL)
    
    def _add_section_21_content_docx(self, doc: Document, params: Dict):
        """Add Section 21 Notice content."""
        self._add_docx_section(doc, "SECTION 21 NOTICE")
        doc.add_paragraph("Housing Act 1988, Section 21(1) or (4)")
        doc.add_paragraph()
        
        doc.add_paragraph(f"Property Address: {params.get('property_address', '[To be completed]')}")
        doc.add_paragraph(f"Tenant Name(s): {params.get('tenant_names', '[To be completed]')}")
        doc.add_paragraph(f"Tenancy Start Date: {params.get('tenancy_start_date', '[To be completed]')}")
        doc.add_paragraph()
        
        doc.add_paragraph(
            "IMPORTANT NOTICE: You are required to leave the above property on or after "
            f"{params.get('notice_expiry_date', '[Date to be calculated]')}."
        )
        doc.add_paragraph()
        
        # Legal text
        legal_para = doc.add_paragraph()
        legal_para.add_run(
            "This notice is given pursuant to Section 21 of the Housing Act 1988. "
            "The landlord requires possession of the dwelling-house after the end of "
            "the fixed term of the tenancy or, if it is a periodic tenancy, "
            "after the end of a period of the tenancy. "
        )
        
        doc.add_paragraph()
        doc.add_paragraph(f"Landlord/Agent: {params.get('landlord_name', '[Name]')}")
        doc.add_paragraph(f"Address: {params.get('landlord_address', '[Address]')}")
        doc.add_paragraph(f"Date: {datetime.now(timezone.utc).strftime('%d %B %Y')}")
        doc.add_paragraph("Signature: _________________________")
    
    def _add_tenancy_content_docx(self, doc: Document, params: Dict):
        """Add Tenancy Agreement content."""
        self._add_docx_section(doc, "ASSURED SHORTHOLD TENANCY AGREEMENT")
        
        # Parties
        doc.add_paragraph("PARTIES TO THIS AGREEMENT:")
        doc.add_paragraph(f"Landlord: {params.get('landlord_name', '[Landlord Name]')}")
        doc.add_paragraph(f"Landlord Address: {params.get('landlord_address', '[Address]')}")
        doc.add_paragraph(f"Tenant(s): {params.get('tenant_names', '[Tenant Names]')}")
        doc.add_paragraph()
        
        # Property
        doc.add_paragraph("PROPERTY:")
        doc.add_paragraph(f"Address: {params.get('property_address', '[Property Address]')}")
        doc.add_paragraph()
        
        # Terms
        doc.add_paragraph("TENANCY TERMS:")
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        terms = [
            ("Term:", f"{params.get('term_months', '12')} months"),
            ("Start Date:", params.get('tenancy_start_date', '[Date]')),
            ("Monthly Rent:", f"£{params.get('monthly_rent', '[Amount]')}"),
            ("Deposit:", f"£{params.get('deposit_amount', '[Amount]')}"),
            ("Payment Date:", params.get('rent_due_day', '1st of each month')),
        ]
        for i, (label, value) in enumerate(terms):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        doc.add_paragraph("SIGNATURES:")
        doc.add_paragraph("Landlord: _________________________  Date: ___________")
        doc.add_paragraph("Tenant: _________________________  Date: ___________")
    
    def _add_audit_content_docx(self, doc: Document, params: Dict):
        """Add Compliance Audit content."""
        self._add_docx_section(doc, "COMPLIANCE AUDIT REPORT")
        
        doc.add_paragraph(f"Property: {params.get('property_address', '[Property Address]')}")
        doc.add_paragraph(f"Audit Date: {datetime.now(timezone.utc).strftime('%d %B %Y')}")
        doc.add_paragraph(f"Property Type: {params.get('property_type', 'Residential')}")
        doc.add_paragraph()
        
        self._add_docx_section(doc, "COMPLIANCE STATUS")
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        headers = table.rows[0].cells
        headers[0].text = "Requirement"
        headers[1].text = "Status"
        headers[2].text = "Expiry Date"
        
        items = [
            ("Gas Safety Certificate", params.get('gas_status', 'To Check'), params.get('gas_expiry', 'N/A')),
            ("EICR", params.get('eicr_status', 'To Check'), params.get('eicr_expiry', 'N/A')),
            ("EPC", params.get('epc_rating', 'To Check'), params.get('epc_expiry', 'N/A')),
            ("Smoke/CO Alarms", params.get('alarms_status', 'To Check'), "Annual Check"),
        ]
        for i, (req, status, expiry) in enumerate(items, 1):
            table.rows[i].cells[0].text = req
            table.rows[i].cells[1].text = status
            table.rows[i].cells[2].text = expiry
        
        doc.add_paragraph()
        self._add_docx_section(doc, "RECOMMENDATIONS")
        doc.add_paragraph("Based on the audit findings, the following actions are recommended:")
        doc.add_paragraph("1. [Recommendation based on findings]")
        doc.add_paragraph("2. [Additional recommendations]")
    
    def _add_market_research_content_docx(self, doc: Document, params: Dict):
        """Add Market Research content."""
        self._add_docx_section(doc, "MARKET RESEARCH REPORT")
        
        doc.add_paragraph(f"Target Area: {params.get('location', '[Location]')}")
        doc.add_paragraph(f"Property Type: {params.get('property_type', 'All Types')}")
        doc.add_paragraph(f"Report Date: {datetime.now(timezone.utc).strftime('%d %B %Y')}")
        doc.add_paragraph()
        
        self._add_docx_section(doc, "MARKET OVERVIEW")
        doc.add_paragraph(
            "This report provides an analysis of the current rental market in the specified area. "
            "The data presented is based on comparable properties and recent market activity."
        )
        
        doc.add_paragraph()
        self._add_docx_section(doc, "RENTAL ESTIMATES")
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        estimates = [
            ("Estimate Type", "Monthly Rent"),
            ("Conservative", f"£{params.get('rent_estimate_low', '800')}"),
            ("Mid-Range", f"£{params.get('rent_estimate_mid', '950')}"),
            ("Optimistic", f"£{params.get('rent_estimate_high', '1100')}"),
        ]
        for i, (label, value) in enumerate(estimates):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            if i == 0:
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    
    def _add_inventory_content_docx(self, doc: Document, params: Dict):
        """Add Inventory Report content."""
        self._add_docx_section(doc, "PROPERTY INVENTORY REPORT")
        
        doc.add_paragraph(f"Property: {params.get('property_address', '[Property Address]')}")
        doc.add_paragraph(f"Inspection Date: {params.get('inspection_date', datetime.now(timezone.utc).strftime('%d %B %Y'))}")
        doc.add_paragraph(f"Report Type: {params.get('checklist_type', 'Move-In')}")
        doc.add_paragraph()
        
        rooms = ["Living Room", "Kitchen", "Bedroom 1", "Bathroom"]
        for room in rooms:
            self._add_docx_section(doc, room.upper())
            table = doc.add_table(rows=4, cols=3)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = "Item"
            table.rows[0].cells[1].text = "Condition"
            table.rows[0].cells[2].text = "Notes"
            
            items = [("Flooring", "Good", ""), ("Walls", "Good", ""), ("Fixtures", "Good", "")]
            for i, (item, condition, notes) in enumerate(items, 1):
                table.rows[i].cells[0].text = item
                table.rows[i].cells[1].text = condition
                table.rows[i].cells[2].text = notes
            doc.add_paragraph()
    
    def _add_general_content_docx(self, doc: Document, params: Dict):
        """Add general document content."""
        self._add_docx_section(doc, "DOCUMENT CONTENT")
        
        if params:
            doc.add_paragraph("Parameters provided:")
            for key, value in params.items():
                doc.add_paragraph(f"• {key}: {value}")
        else:
            doc.add_paragraph("No specific parameters were provided for this document.")
    
    def _add_docx_footer(
        self,
        doc: Document,
        order: Dict[str, Any],
        version: int,
        status: DocumentStatus,
    ):
        """Add document footer."""
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run(
            f"Generated by Pleerity Enterprise Ltd | "
            f"{datetime.now(timezone.utc).strftime('%d %B %Y %H:%M UTC')} | "
            f"v{version} ({status.value})"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(*BRAND_GRAY)
        
        if status != DocumentStatus.FINAL:
            notice_para = doc.add_paragraph()
            notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = notice_para.add_run(
                "This is a DRAFT document and has not been approved. "
                "Do not use for official purposes until marked FINAL."
            )
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    # =========================================================================
    # PDF GENERATION
    # =========================================================================
    
    def _generate_pdf(
        self,
        order: Dict[str, Any],
        version: int,
        doc_type: DocumentType,
        status: DocumentStatus,
        regeneration_notes: Optional[str] = None,
    ) -> io.BytesIO:
        """Generate a professional PDF document."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=20*mm,
            leftMargin=20*mm,
            topMargin=25*mm,
            bottomMargin=20*mm,
        )
        
        # Build story (content elements)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.Color(*[c/255 for c in BRAND_NAVY]),
            alignment=TA_CENTER,
            spaceAfter=12,
        )
        
        section_style = ParagraphStyle(
            'Section',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.Color(*[c/255 for c in BRAND_TEAL]),
            spaceBefore=12,
            spaceAfter=6,
        )
        
        normal_style = styles['Normal']
        
        # Header
        story.append(Paragraph("PLEERITY ENTERPRISE LTD", title_style))
        
        # Status watermark for drafts
        if status != DocumentStatus.FINAL:
            watermark_style = ParagraphStyle(
                'Watermark',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.red,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold',
            )
            story.append(Paragraph(f"*** {status.value} DOCUMENT ***", watermark_style))
        
        story.append(Spacer(1, 12))
        story.append(Paragraph(order.get("service_name", "Document"), title_style))
        
        ref_style = ParagraphStyle(
            'Reference',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.gray,
            alignment=TA_CENTER,
        )
        story.append(Paragraph(
            f"Reference: {order.get('order_id')} | Version: v{version}",
            ref_style
        ))
        
        story.append(HRFlowable(width="100%", thickness=1, color=colors.gray))
        story.append(Spacer(1, 12))
        
        # Customer Details
        story.append(Paragraph("CUSTOMER DETAILS", section_style))
        customer = order.get("customer", {})
        customer_data = [
            ["Full Name:", customer.get("full_name", "Not provided")],
            ["Email:", customer.get("email", "Not provided")],
            ["Phone:", customer.get("phone", "Not provided")],
            ["Company:", customer.get("company", "N/A")],
        ]
        customer_table = Table(customer_data, colWidths=[100, 350])
        customer_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(customer_table)
        story.append(Spacer(1, 12))
        
        # Service Details
        story.append(Paragraph("SERVICE DETAILS", section_style))
        service_data = [
            ["Service:", order.get("service_name", "N/A")],
            ["Service Code:", order.get("service_code", "N/A")],
            ["Category:", order.get("service_category", "N/A")],
        ]
        service_table = Table(service_data, colWidths=[100, 350])
        service_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(service_table)
        story.append(Spacer(1, 12))
        
        # Document-type specific content
        params = order.get("parameters", {})
        if doc_type == DocumentType.SECTION_21_NOTICE:
            self._add_section_21_content_pdf(story, params, section_style, normal_style)
        elif doc_type == DocumentType.TENANCY_AGREEMENT:
            self._add_tenancy_content_pdf(story, params, section_style, normal_style)
        elif doc_type == DocumentType.COMPLIANCE_AUDIT:
            self._add_audit_content_pdf(story, params, section_style, normal_style)
        elif doc_type == DocumentType.MARKET_RESEARCH:
            self._add_market_research_content_pdf(story, params, section_style, normal_style)
        elif doc_type == DocumentType.INVENTORY_REPORT:
            self._add_inventory_content_pdf(story, params, section_style, normal_style)
        else:
            self._add_general_content_pdf(story, params, section_style, normal_style)
        
        # Regeneration notes
        if regeneration_notes:
            story.append(Spacer(1, 12))
            story.append(Paragraph("REGENERATION NOTES", section_style))
            story.append(Paragraph(regeneration_notes, normal_style))
        
        # Footer
        story.append(Spacer(1, 24))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.gray))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.gray,
            alignment=TA_CENTER,
        )
        story.append(Paragraph(
            f"Generated by Pleerity Enterprise Ltd | "
            f"{datetime.now(timezone.utc).strftime('%d %B %Y %H:%M UTC')} | "
            f"v{version} ({status.value})",
            footer_style
        ))
        
        if status != DocumentStatus.FINAL:
            draft_notice = ParagraphStyle(
                'DraftNotice',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.red,
                alignment=TA_CENTER,
            )
            story.append(Paragraph(
                "This is a DRAFT document. Do not use for official purposes until marked FINAL.",
                draft_notice
            ))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _add_section_21_content_pdf(self, story, params, section_style, normal_style):
        """Add Section 21 content to PDF."""
        story.append(Paragraph("SECTION 21 NOTICE", section_style))
        story.append(Paragraph("Housing Act 1988, Section 21(1) or (4)", normal_style))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph(f"Property Address: {params.get('property_address', '[To be completed]')}", normal_style))
        story.append(Paragraph(f"Tenant Name(s): {params.get('tenant_names', '[To be completed]')}", normal_style))
        story.append(Paragraph(f"Tenancy Start Date: {params.get('tenancy_start_date', '[To be completed]')}", normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph(
            f"<b>IMPORTANT NOTICE:</b> You are required to leave the above property on or after "
            f"{params.get('notice_expiry_date', '[Date to be calculated]')}.",
            normal_style
        ))
    
    def _add_tenancy_content_pdf(self, story, params, section_style, normal_style):
        """Add Tenancy Agreement content to PDF."""
        story.append(Paragraph("ASSURED SHORTHOLD TENANCY AGREEMENT", section_style))
        story.append(Spacer(1, 6))
        
        story.append(Paragraph(f"Landlord: {params.get('landlord_name', '[Landlord Name]')}", normal_style))
        story.append(Paragraph(f"Tenant(s): {params.get('tenant_names', '[Tenant Names]')}", normal_style))
        story.append(Paragraph(f"Property: {params.get('property_address', '[Property Address]')}", normal_style))
        story.append(Spacer(1, 12))
        
        terms_data = [
            ["Term:", f"{params.get('term_months', '12')} months"],
            ["Monthly Rent:", f"£{params.get('monthly_rent', '[Amount]')}"],
            ["Deposit:", f"£{params.get('deposit_amount', '[Amount]')}"],
        ]
        terms_table = Table(terms_data, colWidths=[100, 200])
        terms_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('BOX', (0, 0), (-1, -1), 1, colors.black),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(terms_table)
    
    def _add_audit_content_pdf(self, story, params, section_style, normal_style):
        """Add Compliance Audit content to PDF."""
        story.append(Paragraph("COMPLIANCE AUDIT REPORT", section_style))
        story.append(Paragraph(f"Property: {params.get('property_address', '[Property Address]')}", normal_style))
        story.append(Paragraph(f"Audit Date: {datetime.now(timezone.utc).strftime('%d %B %Y')}", normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("COMPLIANCE STATUS", section_style))
        status_data = [
            ["Requirement", "Status", "Expiry"],
            ["Gas Safety", params.get('gas_status', 'To Check'), params.get('gas_expiry', 'N/A')],
            ["EICR", params.get('eicr_status', 'To Check'), params.get('eicr_expiry', 'N/A')],
            ["EPC", params.get('epc_rating', 'To Check'), params.get('epc_expiry', 'N/A')],
            ["Smoke/CO Alarms", params.get('alarms_status', 'To Check'), "Annual"],
        ]
        status_table = Table(status_data, colWidths=[150, 100, 100])
        status_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(*[c/255 for c in BRAND_TEAL])),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('BOX', (0, 0), (-1, -1), 1, colors.black),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(status_table)
    
    def _add_market_research_content_pdf(self, story, params, section_style, normal_style):
        """Add Market Research content to PDF."""
        story.append(Paragraph("MARKET RESEARCH REPORT", section_style))
        story.append(Paragraph(f"Target Area: {params.get('location', '[Location]')}", normal_style))
        story.append(Paragraph(f"Property Type: {params.get('property_type', 'All Types')}", normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("RENTAL ESTIMATES", section_style))
        estimates_data = [
            ["Estimate Type", "Monthly Rent"],
            ["Conservative", f"£{params.get('rent_estimate_low', '800')}"],
            ["Mid-Range", f"£{params.get('rent_estimate_mid', '950')}"],
            ["Optimistic", f"£{params.get('rent_estimate_high', '1100')}"],
        ]
        estimates_table = Table(estimates_data, colWidths=[150, 100])
        estimates_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(*[c/255 for c in BRAND_NAVY])),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('BOX', (0, 0), (-1, -1), 1, colors.black),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
            ('PADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(estimates_table)
    
    def _add_inventory_content_pdf(self, story, params, section_style, normal_style):
        """Add Inventory Report content to PDF."""
        story.append(Paragraph("PROPERTY INVENTORY REPORT", section_style))
        story.append(Paragraph(f"Property: {params.get('property_address', '[Property Address]')}", normal_style))
        story.append(Paragraph(f"Report Type: {params.get('checklist_type', 'Move-In')}", normal_style))
        story.append(Spacer(1, 12))
        
        for room in ["Living Room", "Kitchen", "Bedroom", "Bathroom"]:
            story.append(Paragraph(room.upper(), section_style))
            room_data = [
                ["Item", "Condition", "Notes"],
                ["Flooring", "Good", ""],
                ["Walls", "Good", ""],
                ["Fixtures", "Good", ""],
            ]
            room_table = Table(room_data, colWidths=[100, 80, 200])
            room_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BOX', (0, 0), (-1, -1), 1, colors.black),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.gray),
                ('PADDING', (0, 0), (-1, -1), 4),
            ]))
            story.append(room_table)
            story.append(Spacer(1, 6))
    
    def _add_general_content_pdf(self, story, params, section_style, normal_style):
        """Add general content to PDF."""
        story.append(Paragraph("DOCUMENT CONTENT", section_style))
        
        if params:
            story.append(Paragraph("Parameters provided:", normal_style))
            for key, value in params.items():
                story.append(Paragraph(f"• {key}: {value}", normal_style))
        else:
            story.append(Paragraph("No specific parameters were provided for this document.", normal_style))


# Create singleton instance
real_document_generator = RealDocumentGenerator()
