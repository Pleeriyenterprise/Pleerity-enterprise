"""
Unit tests for Tier 1 document generation guardrails (DOC_GENERATION_AUDIT.md).

- When model returns non-JSON: pipeline returns failure and no render/GridFS write.
- Generic DOCX renderer does not include raw_response/parse_error keys.
- Managed prompt missing {{INPUT_DATA_JSON}} fails early.
"""
import os
import sys
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

BACKEND_DIR = Path(__file__).resolve().parent.parent
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))
os.chdir(BACKEND_DIR)

pytestmark = pytest.mark.asyncio


# -----------------------------------------------------------------------------
# 1) Non-JSON response: pipeline fails, no template_renderer / document_versions_v2
# -----------------------------------------------------------------------------
class TestParseFailureHardFail:
    """When _execute_gpt would return non-JSON, pipeline returns failure and does not call render."""

    @pytest.mark.asyncio
    async def test_non_json_response_returns_failure_and_no_render(self):
        from services.document_orchestrator import (
            document_orchestrator,
            OrchestrationResult,
            OrchestrationStatus,
        )
        from services.template_renderer import template_renderer

        order_id = "guardrail-test-order-1"
        order = {
            "order_id": order_id,
            "service_code": "AI_WF_BLUEPRINT",
            "status": "PAID",
            "order_ref": "GRD-001",
        }
        intake_data = {"business_description": "Test"}

        async def mock_validate(*args, **kwargs):
            return True, "", order

        with (
            patch(
                "services.document_orchestrator.document_orchestrator.validate_order_for_generation",
                side_effect=mock_validate,
            ),
            patch(
                "services.document_orchestrator.prompt_manager_bridge.get_prompt_for_service",
                new_callable=AsyncMock,
            ) as mock_get_prompt,
            patch(
                "services.document_orchestrator.document_orchestrator._execute_gpt",
                new_callable=AsyncMock,
                side_effect=ValueError("LLM output not valid JSON"),
            ),
            patch.object(
                template_renderer,
                "render_from_orchestration",
                new_callable=AsyncMock,
            ) as mock_render,
            patch("services.document_orchestrator.database.get_db") as mock_get_db,
        ):
            from services.gpt_prompt_registry import (
                AI_WF_BLUEPRINT_PROMPT,
                get_prompt_for_service as get_legacy,
            )
            from services.prompt_manager_bridge import ManagedPromptInfo

            # Return legacy prompt (no INPUT_DATA_JSON check) so we reach _execute_gpt
            mock_get_prompt.return_value = (AI_WF_BLUEPRINT_PROMPT, ManagedPromptInfo(
                template_id="legacy",
                version=1,
                service_code="AI_WF_BLUEPRINT",
                doc_type="AI_WF_BLUEPRINT",
                name="Test",
                source="legacy_registry",
            ))
            mock_db = MagicMock()
            mock_db.orders = MagicMock()
            mock_db.orders.find_one = AsyncMock(return_value=order)
            mock_db.orders.update_one = AsyncMock()
            mock_db.__getitem__ = MagicMock(return_value=MagicMock(update_one=AsyncMock()))
            mock_get_db.return_value = mock_db

            result = await document_orchestrator.execute_full_pipeline(
                order_id=order_id,
                intake_data=intake_data,
            )

        assert isinstance(result, OrchestrationResult)
        assert result.success is False
        assert result.error_message == "LLM output not valid JSON"
        assert result.status == OrchestrationStatus.FAILED
        mock_render.assert_not_called()
        # Failure semantics: order must be marked FAILED with error metadata
        update_calls = mock_db.orders.update_one.call_args_list
        assert len(update_calls) >= 1
        set_payload = update_calls[-1][0][1].get("$set", {})
        assert set_payload.get("orchestration_status") == "FAILED"
        assert "last_orchestration_error" in set_payload
        err = set_payload["last_orchestration_error"]
        assert err.get("error_code") == "LLM_INVALID_JSON"
        assert err.get("stage") == "gpt"
        assert "last_orchestration_failed_at" in set_payload


# -----------------------------------------------------------------------------
# 2) Generic DOCX renderer skips raw_response and parse_error
# -----------------------------------------------------------------------------
class TestGenericDocxSkipsErrorKeys:
    """_render_generic_content must not add sections for raw_response or parse_error."""

    def test_generic_content_does_not_include_raw_response_or_parse_error(self):
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from services.template_renderer import TemplateRenderer

        doc = Document()
        styles = doc.styles
        if "CustomHeading" not in [s.name for s in styles]:
            styles.add_style("CustomHeading", WD_STYLE_TYPE.PARAGRAPH)

        renderer = TemplateRenderer()
        output = {
            "raw_response": "raw intake or garbage text",
            "parse_error": "Expecting value: line 1 column 1",
            "data_gaps_flagged": ["ignored"],
            "summary": "Good content only this should appear",
        }
        renderer._render_generic_content(doc, output)

        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Good content only this should appear" in full_text
        assert "raw intake or garbage text" not in full_text
        assert "Expecting value" not in full_text
        assert "Raw Response" not in full_text
        assert "Parse Error" not in full_text


# -----------------------------------------------------------------------------
# 3) Managed prompt missing {{INPUT_DATA_JSON}} fails early
# -----------------------------------------------------------------------------
class TestManagedPromptInputDataJsonGuard:
    """When using managed prompt and template lacks {{INPUT_DATA_JSON}}, fail before calling model."""

    @pytest.mark.asyncio
    async def test_managed_prompt_missing_placeholder_fails_early(self):
        from services.document_orchestrator import (
            document_orchestrator,
            OrchestrationResult,
            OrchestrationStatus,
        )
        from services.gpt_prompt_registry import PromptDefinition, PromptType
        from services.prompt_manager_bridge import ManagedPromptInfo

        order_id = "guardrail-test-order-2"
        order = {
            "order_id": order_id,
            "service_code": "AI_WF_BLUEPRINT",
            "status": "PAID",
            "order_ref": "GRD-002",
        }
        # Managed prompt but template WITHOUT {{INPUT_DATA_JSON}}
        prompt_def = PromptDefinition(
            prompt_id="managed_no_placeholder",
            prompt_type=PromptType.MASTER,
            service_code="AI_WF_BLUEPRINT",
            name="Test",
            description="Test",
            system_prompt="You are a helper.",
            user_prompt_template="No placeholder here - wrong config",
            output_schema={"summary": "string"},
            required_fields=[],
            gpt_sections=[],
        )
        prompt_info = ManagedPromptInfo(
            template_id="tpl-1",
            version=1,
            service_code="AI_WF_BLUEPRINT",
            doc_type="AI_WF_BLUEPRINT",
            name="Test",
            source="prompt_manager",
        )

        async def mock_validate(*args, **kwargs):
            return True, "", order

        with (
            patch(
                "services.document_orchestrator.document_orchestrator.validate_order_for_generation",
                side_effect=mock_validate,
            ),
            patch(
                "services.document_orchestrator.prompt_manager_bridge.get_prompt_for_service",
                new_callable=AsyncMock,
                return_value=(prompt_def, prompt_info),
            ),
            patch(
                "services.document_orchestrator.document_orchestrator._execute_gpt",
                new_callable=AsyncMock,
            ) as mock_gpt,
            patch("services.document_orchestrator.database.get_db") as mock_get_db,
        ):
            mock_db = MagicMock()
            mock_db.orders = MagicMock()
            mock_db.orders.find_one = AsyncMock(return_value=order)
            mock_db.orders.update_one = AsyncMock()
            mock_get_db.return_value = mock_db

            result = await document_orchestrator.execute_full_pipeline(
                order_id=order_id,
                intake_data={"business_description": "Test"},
            )

        assert isinstance(result, OrchestrationResult)
        assert result.success is False
        assert "INPUT_DATA_JSON" in (result.error_message or "")
        assert "configuration error" in (result.error_message or "").lower()
        assert result.status == OrchestrationStatus.FAILED
        mock_gpt.assert_not_called()


# -----------------------------------------------------------------------------
# Failure semantics: orchestration_status FAILED + last_orchestration_error
# -----------------------------------------------------------------------------
class TestAdminGenerateFailurePersistsError:
    """Admin POST /generate failure must persist last_orchestration_error and set orchestration_status FAILED."""

    @pytest.mark.asyncio
    async def test_admin_generate_failure_persists_orchestration_error_and_status(self):
        from services.document_orchestrator import (
            document_orchestrator,
            OrchestrationResult,
            OrchestrationStatus,
        )

        order_id = "fail-sem-order-1"
        failed_result = OrchestrationResult(
            success=False,
            status=OrchestrationStatus.FAILED,
            service_code="AI_WF_BLUEPRINT",
            order_id=order_id,
            error_message="LLM output not valid JSON",
            prompt_version_used={"template_id": "t1", "version": 1},
        )
        mock_db = MagicMock()
        mock_db.orders = MagicMock()
        mock_db.orders.update_one = AsyncMock()
        mock_db.__getitem__ = MagicMock(return_value=MagicMock(update_one=AsyncMock()))

        with patch("services.document_orchestrator.database.get_db", return_value=mock_db):
            await document_orchestrator.finalize_orchestration_failure(
                failed_result, stage="pipeline", error_code="GENERATION_FAILED"
            )

        update_calls = mock_db.orders.update_one.call_args_list
        assert len(update_calls) >= 1
        set_payload = update_calls[0][0][1].get("$set", {})
        assert set_payload.get("orchestration_status") == "FAILED"
        assert "last_orchestration_error" in set_payload
        err = set_payload["last_orchestration_error"]
        assert err.get("error_code") == "GENERATION_FAILED"
        assert err.get("stage") == "pipeline"
        assert err.get("error_message") == "LLM output not valid JSON"
        assert "last_orchestration_failed_at" in set_payload


class TestWF2FailureSetsOrderAndOrchestrationFailed:
    """WF2 generation failure must set order.status FAILED and orchestration_status FAILED."""

    @pytest.mark.asyncio
    async def test_wf2_failure_sets_status_failed_and_orchestration_failed(self):
        from services.workflow_automation_service import WorkflowAutomationService
        from services.document_orchestrator import (
            document_orchestrator,
            OrchestrationResult,
            OrchestrationStatus,
        )
        from services.order_workflow import OrderStatus

        order_id = "wf2-fail-order-1"
        order = {
            "order_id": order_id,
            "service_code": "AI_WF_BLUEPRINT",
            "status": OrderStatus.QUEUED.value,
            "parameters": {"business_description": "Test"},
        }
        exec_id = "wf2-exec-123"
        failed_result = OrchestrationResult(
            success=False,
            status=OrchestrationStatus.FAILED,
            service_code="AI_WF_BLUEPRINT",
            order_id=order_id,
            error_message="LLM output not valid JSON",
            execution_id=exec_id,
        )
        mock_db = MagicMock()
        mock_db.orders = MagicMock()
        mock_db.orders.find_one = AsyncMock(return_value=order)
        mock_db.orders.update_one = AsyncMock()
        mock_db.__getitem__ = MagicMock(return_value=MagicMock(update_one=AsyncMock()))

        transition_calls = []

        async def capture_transition(*args, **kwargs):
            transition_calls.append((args, kwargs))
            return order

        with (
            patch("services.workflow_automation_service.database.get_db", return_value=mock_db),
            patch(
                "services.workflow_automation_service.get_order",
                new_callable=AsyncMock,
                return_value=order,
            ),
            patch(
                "services.workflow_automation_service.transition_order_state",
                side_effect=capture_transition,
            ),
            patch(
                "services.document_orchestrator.database.get_db",
                return_value=mock_db,
            ),
            patch.object(
                document_orchestrator,
                "execute_generation",
                new_callable=AsyncMock,
                return_value=failed_result,
            ),
        ):
            wf = WorkflowAutomationService()
            out = await wf.wf2_queue_to_generation(order_id)

        assert out.get("success") is False
        assert out.get("status") == "FAILED"
        assert len(transition_calls) >= 1
        (_, kwargs) = transition_calls[-1]
        assert kwargs.get("new_status") == OrderStatus.FAILED
        set_payload = mock_db.orders.update_one.call_args_list[-1][0][1].get("$set", {})
        assert set_payload.get("orchestration_status") == "FAILED"
        assert "last_orchestration_error" in set_payload


class TestWF4FailureReturnsToReviewWithOrchestrationFailed:
    """WF4 regeneration failure must return to INTERNAL_REVIEW but set orchestration_status FAILED and store error."""

    @pytest.mark.asyncio
    async def test_wf4_failure_returns_internal_review_with_orchestration_failed(self):
        from services.workflow_automation_service import WorkflowAutomationService
        from services.document_orchestrator import (
            document_orchestrator,
            OrchestrationResult,
            OrchestrationStatus,
        )
        from services.order_workflow import OrderStatus

        order_id = "wf4-fail-order-1"
        order = {
            "order_id": order_id,
            "service_code": "AI_WF_BLUEPRINT",
            "status": OrderStatus.REGENERATING.value,
            "parameters": {"business_description": "Test"},
        }
        exec_id = "wf4-exec-456"
        failed_result = OrchestrationResult(
            success=False,
            status=OrchestrationStatus.FAILED,
            service_code="AI_WF_BLUEPRINT",
            order_id=order_id,
            error_message="Rendering failed",
            execution_id=exec_id,
        )
        mock_db = MagicMock()
        mock_db.orders = MagicMock()
        mock_db.orders.find_one = AsyncMock(return_value=order)
        mock_db.orders.update_one = AsyncMock()
        mock_db.__getitem__ = MagicMock(return_value=MagicMock(update_one=AsyncMock()))

        transition_calls = []

        async def capture_transition(*args, **kwargs):
            transition_calls.append((args, kwargs))
            return order

        with (
            patch("services.workflow_automation_service.database.get_db", return_value=mock_db),
            patch(
                "services.workflow_automation_service.get_order",
                new_callable=AsyncMock,
                return_value=order,
            ),
            patch(
                "services.workflow_automation_service.transition_order_state",
                side_effect=capture_transition,
            ),
            patch(
                "services.document_orchestrator.database.get_db",
                return_value=mock_db,
            ),
            patch.object(
                document_orchestrator,
                "execute_generation",
                new_callable=AsyncMock,
                return_value=failed_result,
            ),
        ):
            wf = WorkflowAutomationService()
            out = await wf.wf4_regeneration(order_id, "Please fix the summary section.")

        assert out.get("success") is False
        assert out.get("status") == "INTERNAL_REVIEW"
        assert len(transition_calls) >= 1
        (_, kwargs) = transition_calls[-1]
        assert kwargs.get("new_status") == OrderStatus.INTERNAL_REVIEW
        set_payload = mock_db.orders.update_one.call_args_list[-1][0][1].get("$set", {})
        assert set_payload.get("orchestration_status") == "FAILED"
        assert "last_orchestration_error" in set_payload


class TestSingleFailedRunProducesOneExecutionRecord:
    """A single failed run (orchestrator + finalize) must produce exactly one FAILED execution record."""

    @pytest.mark.asyncio
    async def test_only_one_failed_execution_record_per_run(self):
        from services.document_orchestrator import (
            document_orchestrator,
            OrchestrationResult,
            OrchestrationStatus,
        )
        from services.gpt_prompt_registry import AI_WF_BLUEPRINT_PROMPT
        from services.prompt_manager_bridge import ManagedPromptInfo

        order_id = "single-fail-run-1"
        order = {
            "order_id": order_id,
            "service_code": "AI_WF_BLUEPRINT",
            "status": "PAID",
            "order_ref": "SFR-001",
        }
        failed_execution_records = []

        async def capture_upsert(*args, **kwargs):
            filter_arg = args[0] if args else {}
            update_arg = args[1] if len(args) > 1 else {}
            if kwargs.get("upsert") and filter_arg.get("status") == "FAILED":
                eid = filter_arg.get("execution_id")
                if eid and not any(r.get("execution_id") == eid for r in failed_execution_records):
                    doc = {**filter_arg, **(update_arg.get("$set") or {}), **(update_arg.get("$setOnInsert") or {})}
                    failed_execution_records.append(doc)

        mock_coll = MagicMock()
        mock_coll.update_one = AsyncMock(side_effect=capture_upsert)
        mock_db = MagicMock()
        mock_db.orders = MagicMock()
        mock_db.orders.find_one = AsyncMock(return_value=order)
        mock_db.orders.update_one = AsyncMock()
        mock_db.__getitem__ = MagicMock(return_value=mock_coll)

        async def mock_validate(*args, **kwargs):
            return True, "", order

        with (
            patch(
                "services.document_orchestrator.document_orchestrator.validate_order_for_generation",
                side_effect=mock_validate,
            ),
            patch(
                "services.document_orchestrator.prompt_manager_bridge.get_prompt_for_service",
                new_callable=AsyncMock,
                return_value=(
                    AI_WF_BLUEPRINT_PROMPT,
                    ManagedPromptInfo(
                        template_id="legacy",
                        version=1,
                        service_code="AI_WF_BLUEPRINT",
                        doc_type="AI_WF_BLUEPRINT",
                        name="Test",
                        source="legacy_registry",
                    ),
                ),
            ),
            patch(
                "services.document_orchestrator.document_orchestrator._execute_gpt",
                new_callable=AsyncMock,
                side_effect=ValueError("LLM output not valid JSON"),
            ),
            patch("services.document_orchestrator.database.get_db", return_value=mock_db),
        ):
            result = await document_orchestrator.execute_full_pipeline(
                order_id=order_id,
                intake_data={"business_description": "Test"},
            )
        assert result.success is False
        assert result.execution_id is not None
        exec_id = result.execution_id
        await document_orchestrator.finalize_orchestration_failure(
            result, stage="pipeline", error_code="GENERATION_FAILED"
        )
        assert len(failed_execution_records) == 1, "expected exactly one FAILED execution record for one failed run"
        assert failed_execution_records[0].get("execution_id") == exec_id
        assert failed_execution_records[0].get("status") == "FAILED"
        assert failed_execution_records[0].get("order_id") == order_id

    @pytest.mark.asyncio
    async def test_error_message_truncated_to_max_length(self):
        from services.document_orchestrator import (
            document_orchestrator,
            OrchestrationResult,
            OrchestrationStatus,
            MAX_ERROR_MESSAGE_LENGTH,
        )

        long_message = "x" * (MAX_ERROR_MESSAGE_LENGTH + 500)
        failed_result = OrchestrationResult(
            success=False,
            status=OrchestrationStatus.FAILED,
            service_code="AI_WF_BLUEPRINT",
            order_id="trunc-order-1",
            error_message=long_message,
            execution_id="trunc-exec-1",
        )
        mock_db = MagicMock()
        mock_db.orders = MagicMock()
        mock_db.orders.update_one = AsyncMock()
        mock_db.__getitem__ = MagicMock(return_value=MagicMock(update_one=AsyncMock()))

        with patch("services.document_orchestrator.database.get_db", return_value=mock_db):
            await document_orchestrator.finalize_orchestration_failure(
                failed_result, stage="pipeline", error_code="TEST"
            )

        set_payload = mock_db.orders.update_one.call_args[0][1].get("$set", {})
        stored_msg = set_payload.get("last_orchestration_error", {}).get("error_message", "")
        assert len(stored_msg) == MAX_ERROR_MESSAGE_LENGTH
        assert stored_msg == "x" * MAX_ERROR_MESSAGE_LENGTH


# -----------------------------------------------------------------------------
# Tier 2.1: Idempotency for document generation runs
# -----------------------------------------------------------------------------
class TestIdempotencyTwoIdenticalCallsOneVersion:
    """Two identical calls (same idempotency_key) must produce only one new document version."""

    @pytest.mark.asyncio
    async def test_two_identical_calls_produce_only_one_render(self):
        from services.document_orchestrator import document_orchestrator, OrchestrationStatus
        from services.template_renderer import template_renderer
        from services.gpt_prompt_registry import AI_WF_BLUEPRINT_PROMPT
        from services.prompt_manager_bridge import ManagedPromptInfo

        order_id = "idem-order-1"
        order = {
            "order_id": order_id,
            "service_code": "AI_WF_BLUEPRINT",
            "status": "IN_PROGRESS",
            "order_ref": "IDEM-001",
        }
        executions_stored = []

        async def insert_one_side_effect(doc):
            executions_stored.append(doc)

        async def find_one_side_effect(filter_arg, sort=None):
            key = (filter_arg or {}).get("idempotency_key")
            if key and executions_stored:
                for d in reversed(executions_stored):
                    if d.get("idempotency_key") == key:
                        return d
            return None

        mock_coll = MagicMock()
        mock_coll.insert_one = AsyncMock(side_effect=insert_one_side_effect)
        mock_coll.find_one = AsyncMock(side_effect=find_one_side_effect)
        mock_db = MagicMock()
        mock_db.orders = MagicMock()
        mock_db.orders.find_one = AsyncMock(return_value=order)
        mock_db.orders.update_one = AsyncMock()
        mock_db.__getitem__ = MagicMock(return_value=mock_coll)

        async def mock_validate(*args, **kwargs):
            return True, "", order

        valid_output = {"executive_summary": "Summary", "sections": []}
        from services.template_renderer import (
            RenderResult,
            RenderStatus,
            RenderedDocument,
        )
        render_return = RenderResult(
            success=True,
            order_id=order_id,
            version=1,
            status=RenderStatus.REVIEW_PENDING,
            docx=RenderedDocument("a.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", b"", "h", 1, "docx"),
            pdf=RenderedDocument("a.pdf", "application/pdf", b"", "h", 1, "pdf"),
            json_output_hash="j",
            render_time_ms=0,
            error_message=None,
        )

        with (
            patch(
                "services.document_orchestrator.document_orchestrator.validate_order_for_generation",
                side_effect=mock_validate,
            ),
            patch(
                "services.document_orchestrator.prompt_manager_bridge.get_prompt_for_service",
                new_callable=AsyncMock,
                return_value=(
                    AI_WF_BLUEPRINT_PROMPT,
                    ManagedPromptInfo(
                        template_id="t1",
                        version=1,
                        service_code="AI_WF_BLUEPRINT",
                        doc_type="AI_WF_BLUEPRINT",
                        name="Test",
                        source="legacy_registry",
                    ),
                ),
            ),
            patch(
                "services.document_orchestrator.document_orchestrator._execute_gpt",
                new_callable=AsyncMock,
                return_value=(valid_output, {"prompt_tokens": 0, "completion_tokens": 0}),
            ),
            patch("services.document_orchestrator.database.get_db", return_value=mock_db),
            patch.object(
                template_renderer,
                "render_from_orchestration",
                new_callable=AsyncMock,
                return_value=render_return,
            ),
        ):
            r1 = await document_orchestrator.execute_full_pipeline(
                order_id=order_id,
                intake_data={"business_description": "Test"},
                regeneration=False,
                force=False,
            )
            r2 = await document_orchestrator.execute_full_pipeline(
                order_id=order_id,
                intake_data={"business_description": "Test"},
                regeneration=False,
                force=False,
            )
        assert r1.success is True
        assert r2.success is True
        assert r1.version == 1
        assert r2.version == 1
        assert template_renderer.render_from_orchestration.await_count == 1
        assert len(executions_stored) == 1
        assert executions_stored[0].get("idempotency_key")


class TestIdempotencyForceRetryAfterFailed:
    """force=true must create a new run even when previous run with same key failed."""

    @pytest.mark.asyncio
    async def test_force_true_allows_rerun_after_failed(self):
        from services.document_orchestrator import document_orchestrator, OrchestrationStatus
        from services.gpt_prompt_registry import AI_WF_BLUEPRINT_PROMPT
        from services.prompt_manager_bridge import ManagedPromptInfo

        order_id = "force-order-1"
        order = {"order_id": order_id, "service_code": "AI_WF_BLUEPRINT", "status": "IN_PROGRESS"}
        existing_failed = {
            "order_id": order_id,
            "service_code": "AI_WF_BLUEPRINT",
            "status": "FAILED",
            "idempotency_key": "force-order-1|AI_WF_BLUEPRINT|GEN|IN_PROGRESS|t1:1|_",
            "execution_id": "e1",
        }
        mock_coll = MagicMock()
        mock_coll.find_one = AsyncMock(return_value=existing_failed)
        mock_coll.insert_one = AsyncMock()
        mock_coll.update_one = AsyncMock()
        mock_db = MagicMock()
        mock_db.orders = MagicMock()
        mock_db.orders.find_one = AsyncMock(return_value=order)
        mock_db.orders.update_one = AsyncMock()
        mock_db.__getitem__ = MagicMock(return_value=mock_coll)

        async def mock_validate(*args, **kwargs):
            return True, "", order

        with (
            patch(
                "services.document_orchestrator.document_orchestrator.validate_order_for_generation",
                side_effect=mock_validate,
            ),
            patch(
                "services.document_orchestrator.prompt_manager_bridge.get_prompt_for_service",
                new_callable=AsyncMock,
                return_value=(
                    AI_WF_BLUEPRINT_PROMPT,
                    ManagedPromptInfo("t1", 1, "AI_WF_BLUEPRINT", "AI_WF_BLUEPRINT", "Test", "legacy_registry"),
                ),
            ),
            patch(
                "services.document_orchestrator.document_orchestrator._execute_gpt",
                new_callable=AsyncMock,
                side_effect=ValueError("LLM output not valid JSON"),
            ),
            patch("services.document_orchestrator.database.get_db", return_value=mock_db),
        ):
            result = await document_orchestrator.execute_full_pipeline(
                order_id=order_id,
                intake_data={"x": "y"},
                regeneration=False,
                force=True,
            )
        assert result.success is False
        assert "LLM output not valid JSON" in (result.error_message or "")
        assert mock_coll.update_one.called or mock_coll.insert_one.called


class TestIdempotencyKeyNoPII:
    """idempotency_key must not contain raw PII (e.g. regeneration_notes are hashed)."""

    def test_idempotency_key_hashes_notes_no_raw_pii(self):
        from services.document_orchestrator import _compute_idempotency_key

        key = _compute_idempotency_key(
            order_id="ord-1",
            service_code="AI_WF_BLUEPRINT",
            regeneration=True,
            order_status="INTERNAL_REVIEW",
            prompt_version_used={"template_id": "t1", "version": 1},
            regeneration_notes="John Smith, 123 Main St, London. Please fix the summary.",
        )
        assert "John" not in key
        assert "Smith" not in key
        assert "123 Main" not in key
        assert "London" not in key
        assert "ord-1" in key
        assert "AI_WF_BLUEPRINT" in key
        assert "REGEN" in key
        assert "t1" in key
        assert "1" in key
        import re
        assert re.search(r"[a-f0-9]{16}", key)
